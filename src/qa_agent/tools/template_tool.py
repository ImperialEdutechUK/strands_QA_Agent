"""Template document interpretation.

A QA template can arrive as:
  * an image (PNG / JPEG / WebP / etc) — OCR'd directly;
  * a PDF — every page's text is extracted, and every embedded image on every
    page is OCR'd, so screenshots-of-rules pasted into a brief still surface;
  * a Word .docx — every paragraph + table cell is read, and every embedded
    image is OCR'd.

In every case the extracted text is merged into a single string and handed to
the LLM, which produces the structured rule list (`{summary, rules: [...]}`).
"""

from __future__ import annotations

import io
import logging
import os
import re
from pathlib import Path

import pytesseract
from PIL import Image

from ..llm_client import call_llm_json
from ..security import (
    ALLOWED_DOC_SUFFIXES,
    ALLOWED_IMAGE_SUFFIXES,
    UnsafePathError,
    safe_resolve_template,
    truncate_text,
)

logger = logging.getLogger(__name__)

if os.environ.get("TESSERACT_CMD"):
    pytesseract.pytesseract.tesseract_cmd = os.environ["TESSERACT_CMD"]

SYSTEM = (
    "You convert QA template documents (often supplied as PDFs, Word files, "
    "or images) into a structured rule set that another QA agent can apply "
    "against a course web page. The template is typically a multi-step "
    "checklist; turn EVERY individual check item into its own rule so nothing "
    "is dropped. Keep the original step/section name as a prefix in each rule "
    "so the reviewer can trace it back."
)

# Sections of the checklist that must NOT become rules (the agent has no way to
# run them, or the user explicitly excluded them).
_EXCLUDE_RULE_RE = re.compile(r"winston|ai percentage|human score|readability score", re.I)

# Items that can only be confirmed against the official Qualification
# Specification (not from the page alone). The agent resolves the spec via a
# web search and these are checked against it.
_SPEC_RULE_RE = re.compile(
    r"match(?:es)?\b.*\b(?:specification|spec\b)|qualification (?:number|specification)|"
    r"credit equivalen|glh|tqt|accreditation status|awarding body|"
    r"entry requirement|access duration",
    re.I,
)

SCHEMA_INSTRUCTION = """Return a JSON object with this shape:
{
  "summary": "<one-sentence summary of the template>",
  "rules": [
    {
      "id": "R1",
      "category": "Content" | "Structure" | "Style" | "Accessibility" | "Branding" | "Other",
      "rule": "<the rule, phrased as a check, prefixed with its checklist step/section>",
      "severity": "Critical" | "Minor" | "Info",
      "needs_spec": true | false
    }
  ]
}
Rules:
  * Create one rule per check item in the template — do not merge or summarise
    multiple checklist items into a single rule.
  * Set "needs_spec" to true when the item can only be verified by comparing the
    page against the official Qualification Specification (e.g. "Course Name
    matches the specification", qualification number, level, credit equivalency,
    GLH/TQT, accreditation status, awarding body, entry requirements, access
    duration). Otherwise set it to false.
  * Do NOT create any rule for the "Winston AI Evaluation" section — skip it.
Output ONLY the JSON object."""


# Baseline rules: the SLC / Imperial Learning "Course Page QA Checklist"
# (10-step Front-End QA review), applied to every QA run regardless of any
# template the user also supplies. Each check item in the master checklist is
# its own rule so nothing is silently dropped. IDs use the checklist step as a
# prefix (S01-… etc.) so a reviewer can trace each finding back to the step.
#
# `needs_spec: True` marks items that can only be confirmed against the official
# Qualification / Centre Specification (resolved via the spec_lookup tool);
# everything else is verifiable from the page itself.
#
# Step 10's Winston AI / AI-percentage / Human-Score / readability items are
# deliberately NOT encoded here — they need a third-party tool and are filtered
# by `_EXCLUDE_RULE_RE` as a safety net.
BASELINE_RULES: list[dict] = [
    # ---- Step 01 — General Information (spec match) --------------------------
    {"id": "S01-1", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: the Course Name must match the official Qualification / Centre Specification exactly."},
    {"id": "S01-2", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: the Level (e.g. Level 4, Level 5) must match the specification."},
    {"id": "S01-3", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: the Accreditation Status must match the specification."},
    {"id": "S01-4", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: the Credit Equivalency must match the specification."},
    {"id": "S01-5", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: the Qualification Number must match the specification."},
    {"id": "S01-6", "category": "Content", "severity": "Critical", "needs_spec": False,
     "rule": "Step 01 General Information: Access Duration is correct — the standard is 1 year; Extended Diplomas are 2 years."},
    {"id": "S01-7", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 01 General Information: Guided Learning Hours (GLH) and Total Qualification Time (TQT) match the spec and are not mixed up from another course."},
    {"id": "S01-8", "category": "Content", "severity": "Critical", "needs_spec": False,
     "rule": "Step 01 Content Document Alignment: live page text matches the approved Content Document — no text missed, mispasted, or duplicated."},

    # ---- Step 02 — Header & Course Overview ---------------------------------
    {"id": "S02-1", "category": "Branding", "severity": "Minor", "needs_spec": False,
     "rule": "Step 02 Header & Overview: the awarding body name (e.g. Qualifi) is shown in bold in the course overview."},
    {"id": "S02-2", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 02 Header & Overview: the Course Overview content is present and relevant to this course."},
    {"id": "S02-3", "category": "Structure", "severity": "Minor", "needs_spec": False,
     "rule": "Step 02 Contact Details: a prominent heading sits directly above the WhatsApp call and email contact details."},

    # ---- Step 03 — "What Will I Learn?" & Key Highlights --------------------
    {"id": "S03-1", "category": "Style", "severity": "Minor", "needs_spec": False,
     "rule": "Step 03 What Will I Learn / Key Highlights: sentences are not overly long."},
    {"id": "S03-2", "category": "Structure", "severity": "Minor", "needs_spec": False,
     "rule": "Step 03 What Will I Learn / Key Highlights: alignment is clean and even."},

    # ---- Step 04 — Course Content / Content Table --------------------------
    {"id": "S04-1", "category": "Content", "severity": "Critical", "needs_spec": True,
     "rule": "Step 04 Highlights & Outcomes: Course Highlights match the key parameters of the qualification specification."},
    {"id": "S04-2", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Learning Outcomes: between 4 and 6 outcomes are listed; if too few, request more from the content writer."},
    {"id": "S04-3", "category": "Style", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Learning Outcomes: every learning outcome starts with a verb."},
    {"id": "S04-4", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Who Is This Course For: clear messaging such as 'This course is ideal for…' with a list of suitable individuals."},
    {"id": "S04-5", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Who Is This Course For: misplaced completion statements are removed (e.g. 'On successful completion you will be ready to take…')."},
    {"id": "S04-6", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Certification: the section clearly states the name of the certificate."},
    {"id": "S04-7", "category": "Content", "severity": "Minor", "needs_spec": True,
     "rule": "Step 04 Entry Requirements: match the entry requirements in the Qualification Specification."},
    {"id": "S04-8", "category": "Structure", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 SLC only: the two boxes (blue and yellow) following the content are present."},
    {"id": "S04-9", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Average Completion Timeframe matches the course type — Award: 2–4 months; Certificate: 4–6 months; Diploma: 6–9 months; Extended Diploma: 9–12 months (adjusted where fast-track information has been received)."},
    {"id": "S04-10", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Method of Assessment: assignment-based courses clearly state that no exams are involved."},
    {"id": "S04-11", "category": "Structure", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Method of Assessment: a clear heading titled 'Assessment Overview' sits above the assessment details, and unnecessary detail is removed."},
    {"id": "S04-12", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Method of Assessment: the SLC final line is present — 'Mentor Guidance and Support — Our experienced mentors will guide and support you throughout the assessment process. You can reach them via email.' followed by a yellow box."},
    {"id": "S04-13", "category": "Structure", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Academic Progression Pathway: the header matches 'Next step for a potential academic progression pathway', steps progress logically upward (e.g. Level 5 Diploma → Bachelor's → Master's), and any 'Alternative Pathways' listed are accurate."},
    {"id": "S04-14", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Career Progression: details are accurate, at least 5 careers are listed from highest to lowest salary, and the salaries are up to date."},
    {"id": "S04-15", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 04 Regulatory & Awarding Bodies: regulatory body information and awarding body information are present and up to date."},

    # ---- Step 05 — Course Curriculum ---------------------------------------
    {"id": "S05-1", "category": "Other", "severity": "Critical", "needs_spec": False,
     "rule": "Step 05 Curriculum Code Check: the module breakdown is free of stray website code next to module text (trailing brackets, unformatted semicolons, double question marks, etc.)."},

    # ---- Step 06 — FAQ Section ---------------------------------------------
    {"id": "S06-1", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 06 FAQ: answers strictly match their FAQ questions (FAQ relevance)."},
    {"id": "S06-2", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 06 FAQ: there are no repeating / duplicate FAQs."},
    {"id": "S06-3", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 06 FAQ Volume: approximately 20 FAQs are present; if below volume, flag it so the content writer can generate more."},
    {"id": "S06-4", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 06 FAQ: the first FAQ is 'What is this [Course Name]?'."},

    # ---- Step 07 — Reviews --------------------------------------------------
    {"id": "S07-1", "category": "Content", "severity": "Minor", "needs_spec": False,
     "rule": "Step 07 Reviews: all review dates are real and logical (no impossible dates like Feb 31, no accidental future dates like Jul 2035), and reviews are accurate."},

    # ---- Step 08 — Course Pricing ------------------------------------------
    {"id": "S08-1", "category": "Content", "severity": "Critical", "needs_spec": False,
     "rule": "Step 08 Pricing: course pricing has been added."},

    # ---- Step 09 — Additional Information (Presentation & Consistency) ------
    {"id": "S09-1", "category": "Style", "severity": "Minor", "needs_spec": False,
     "rule": "Step 09 Presentation: alignments are correct throughout the page."},
    {"id": "S09-2", "category": "Other", "severity": "Critical", "needs_spec": False,
     "rule": "Step 09 Presentation: there are no unnecessary coding parts left in the text."},
    {"id": "S09-3", "category": "Style", "severity": "Critical", "needs_spec": False,
     "rule": "Step 09 Presentation: font types and styles are consistent across the page."},
    {"id": "S09-4", "category": "Style", "severity": "Minor", "needs_spec": False,
     "rule": "Step 09 Presentation: bullet points are properly applied and bolding is correctly done."},
    {"id": "S09-5", "category": "Branding", "severity": "Minor", "needs_spec": False,
     "rule": "Step 09 Presentation: awarding body logos are present and not blurred."},
    {"id": "S09-6", "category": "Style", "severity": "Minor", "needs_spec": False,
     "rule": "Step 09 Presentation: capitalisation is appropriate and the correct course name is used properly throughout."},
]


def _clean_rules(rules: list[dict]) -> list[dict]:
    """Drop excluded sections (e.g. Winston AI) and normalise the needs_spec flag.

    The LLM is told to skip the Winston AI Evaluation section, but we filter
    here too as a safety net so an excluded item can never leak into the rule
    set the QA agent runs.
    """
    out: list[dict] = []
    for r in rules or []:
        if not isinstance(r, dict):
            continue
        text = str(r.get("rule", ""))
        if _EXCLUDE_RULE_RE.search(text):
            continue
        r = dict(r)
        # Trust the model's flag if present; otherwise infer from the rule text.
        if "needs_spec" not in r:
            r["needs_spec"] = bool(_SPEC_RULE_RE.search(text))
        else:
            r["needs_spec"] = bool(r["needs_spec"])
        out.append(r)
    return out


def _merge_baseline(rules: list[dict]) -> list[dict]:
    """Prepend the master-checklist baseline rules, dropping any LLM-emitted rule
    that reuses a baseline ID. Each baseline rule keeps its own needs_spec flag
    (Step 01 / spec-match items are True; the rest False)."""
    baseline_ids = {r["id"] for r in BASELINE_RULES}
    cleaned = _clean_rules(rules)
    extra = [r for r in cleaned if r.get("id") not in baseline_ids]
    baseline = [dict(r, needs_spec=bool(r.get("needs_spec", False))) for r in BASELINE_RULES]
    return baseline + extra


# ---------------------------------------------------------------------------
# Per-format extraction helpers
# ---------------------------------------------------------------------------

def _ocr_image(img: Image.Image) -> str:
    try:
        return (pytesseract.image_to_string(img) or "").strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("OCR failed on embedded image: %s", exc)
        return ""


def _extract_image(path: Path) -> str:
    with Image.open(path) as img:
        return _ocr_image(img)


def _extract_pdf(path: Path) -> str:
    """Pull text + OCR embedded images out of every page of the PDF."""
    # PyMuPDF imports as `fitz`. Imported lazily so an absent dep only breaks
    # PDF templates, not the whole tool.
    import fitz  # type: ignore

    chunks: list[str] = []
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = (page.get_text("text") or "").strip()
            if text:
                chunks.append(f"[page {page_index} text]\n{text}")

            # Embedded images — useful when a brief is largely screenshots.
            for img_idx, img_info in enumerate(page.get_images(full=True), start=1):
                xref = img_info[0]
                try:
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha >= 4:
                        # CMYK / unsupported colourspace — convert to RGB.
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    png_bytes = pix.tobytes("png")
                    pix = None  # release immediately
                except Exception as exc:  # noqa: BLE001
                    logger.debug("PDF image extract failed (page %s, img %s): %s",
                                 page_index, img_idx, exc)
                    continue
                with Image.open(io.BytesIO(png_bytes)) as img:
                    ocr = _ocr_image(img)
                if ocr:
                    chunks.append(f"[page {page_index} image {img_idx} OCR]\n{ocr}")
    return "\n\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    """Pull paragraphs, table cells, and embedded image OCR out of a DOCX."""
    from docx import Document  # type: ignore

    document = Document(str(path))
    chunks: list[str] = []

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    if paragraphs:
        chunks.append("[paragraphs]\n" + "\n".join(paragraphs))

    table_lines: list[str] = []
    for t_idx, table in enumerate(document.tables, start=1):
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                table_lines.append(" | ".join(cells))
    if table_lines:
        chunks.append("[tables]\n" + "\n".join(table_lines))

    # Inline / floating images live in the package's media parts.
    media_count = 0
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
        except Exception as exc:  # noqa: BLE001
            logger.debug("DOCX image extract failed: %s", exc)
            continue
        try:
            with Image.open(io.BytesIO(blob)) as img:
                ocr = _ocr_image(img)
        except Exception as exc:  # noqa: BLE001
            logger.debug("DOCX image OCR failed: %s", exc)
            continue
        if ocr:
            media_count += 1
            chunks.append(f"[image {media_count} OCR]\n{ocr}")

    return "\n\n".join(chunks).strip()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _baseline_only(reason: str) -> dict:
    """The standard QA checklist with no template-specific extras.

    Used whenever the template can't be turned into extra rules (empty, or the
    LLM call failed). The 41-item master checklist is hard-coded and needs no
    LLM, so the compliance/spec flow still runs against it — a rate-limited
    template parse must NOT wipe out the whole QA run.
    """
    return {"summary": reason, "rules": _merge_baseline([])}


def _llm_rules_from_text(label: str, extracted: str) -> dict:
    if not extracted:
        return _baseline_only(
            f"Empty template ({label}: no text extracted) — standard QA checklist applied."
        )
    prompt = (
        f"{SCHEMA_INSTRUCTION}\n\n"
        f"EXTRACTED CONTENT FROM TEMPLATE ({label}):\n"
        f'"""{truncate_text(extracted)}"""'
    )
    try:
        result = call_llm_json(prompt, system=SYSTEM)
    except Exception as exc:  # noqa: BLE001 - never let a template parse failure
        # sink the whole run; fall back to the standard checklist.
        logger.warning(
            "template parse via LLM failed (%s: %s) — applying the standard QA "
            "checklist only", type(exc).__name__, str(exc)[:160],
        )
        return _baseline_only(
            f"Template could not be parsed ({type(exc).__name__}); the standard "
            "QA checklist was applied instead."
        )
    return {
        "summary": result.get("summary", ""),
        "rules": _merge_baseline(result.get("rules") or []),
    }


def analyse_template(template_path: str) -> dict:
    """Interpret a template document (image / PDF / DOCX) into a rule list.

    Only filesystem paths under the configured allowed roots are accepted; raw
    base64 / data URIs are rejected on purpose to avoid being a generic OCR
    endpoint that processes attacker-supplied bytes.

    The raw extracted document text is intentionally NOT returned: the agent
    only needs `summary` + `rules` to drive compliance, and echoing the full
    document body back into Strands' tool-result history bloats every
    subsequent LLM turn and was a primary cause of mid-stream "Network
    connection lost" errors on OpenRouter.
    """
    safe_path = safe_resolve_template(template_path)
    suffix = safe_path.suffix.lower()
    if suffix in ALLOWED_IMAGE_SUFFIXES:
        extracted = _extract_image(safe_path)
        label = "image OCR"
    elif suffix == ".pdf":
        extracted = _extract_pdf(safe_path)
        label = "PDF"
    elif suffix == ".docx":
        extracted = _extract_docx(safe_path)
        label = "Word docx"
    else:
        raise UnsafePathError(f"Unsupported template extension: {suffix}")

    out = _llm_rules_from_text(label, extracted.strip())
    out["source_kind"] = label
    return out


def analyse_template_text(text: str) -> dict:
    text = (text or "").strip()
    if not text:
        return _baseline_only("Empty template — standard QA checklist applied.")
    prompt = f'{SCHEMA_INSTRUCTION}\n\nTEMPLATE TEXT:\n"""{truncate_text(text)}"""'
    try:
        result = call_llm_json(prompt, system=SYSTEM)
    except Exception as exc:  # noqa: BLE001 - keep the run alive on LLM failure
        logger.warning(
            "template text parse via LLM failed (%s: %s) — applying the standard "
            "QA checklist only", type(exc).__name__, str(exc)[:160],
        )
        return _baseline_only(
            f"Template could not be parsed ({type(exc).__name__}); the standard "
            "QA checklist was applied instead."
        )
    return {
        "summary": result.get("summary", ""),
        "rules": _merge_baseline(result.get("rules") or []),
    }
